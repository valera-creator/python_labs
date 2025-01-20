# pip install python-docx

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt, RGBColor


def main():
    # создание пустого документа
    document = Document()

    heading = document.add_heading('Кароче я', level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p1 = document.add_paragraph('учился на курсах кодить, щас только лишь вспоминаю, дальше лень писать')
    # выравнивание
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.bold = True

    p2 = document.add_paragraph('а тут я хочу закодить другое форматирование, ведь так написано в тзшке задания')
    # добавляем отступ слева
    p2.paragraph_format.left_indent = Mm(50)
    # добавляем отступ справа
    p2.paragraph_format.right_indent = Mm(-30)

    p3 = document.add_paragraph()
    p3.add_run("\n")
    p3_run = p3.add_run('а после этого я хочу вспомнить добавление картинки')
    p3_run.font.name = 'Arial'
    p3_run.font.size = Pt(24)
    p3_run.font.color.rgb = RGBColor(0, 0, 255)

    # добавляем и форматируем изображение
    picture_paragraph = document.add_paragraph()
    picture_run = picture_paragraph.add_run()
    picture_run.add_picture('yl.jpg', width=Mm(100), height=Mm(70))
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    document.save('valera.docx')


if __name__ == "__main__":
    main()
