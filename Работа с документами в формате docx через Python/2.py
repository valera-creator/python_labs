# pip install docxtpl
from docxtpl import DocxTemplate
import csv
import os
from docx import Document


def main():
    try:
        with open('data_marathon.csv') as csvfile:
            data = list(csv.reader(csvfile, delimiter=',', quotechar='"'))

    except FileNotFoundError:
        print('нет файла')
        return

    if not data:
        print('пустой файл')
        return

    data = sorted(data, key=lambda x: x[0])

    group_year = {}
    for elem in data:
        if elem[0] not in group_year:
            group_year[elem[0]] = []
        group_year[elem[0]].append(elem)

    doc = DocxTemplate("template.docx")
    main_doc = Document()

    for year, marathons in group_year.items():
        context = {
            "year": year,
            "marathons": [
                {
                    "city": marathon[5],
                    "name": marathon[1],
                    "gender": marathon[2],
                    "time": marathon[4],
                }
                for marathon in marathons
            ]
        }
        doc.render(context)
        doc.save('save_info.docx')

        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)

        for elem in text:
            main_doc.add_paragraph(elem)
        main_doc.add_page_break()
    main_doc.save('res.docx')
    os.remove('save_info.docx')


if __name__ == "__main__":
    main()
